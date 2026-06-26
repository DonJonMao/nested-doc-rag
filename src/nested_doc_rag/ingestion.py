from __future__ import annotations

import hashlib
import json
import time
import uuid
from collections import Counter
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from qdrant_client import QdrantClient, models

from nested_doc_rag.config import AppConfig
from nested_doc_rag.embedding import EmbeddingClient
from nested_doc_rag.evidence_images import (
    attachment_from_registry,
    materialize_xlsx_dispimg_registry,
    registry_by_sheet_row,
)
from nested_doc_rag.io import display_text, write_json, write_jsonl
from nested_doc_rag.retrieval.qdrant_client import build_qdrant_client

SUPPORTED_SUFFIXES = {".xlsx", ".xlsm", ".docx", ".txt", ".md", ".csv"}
MAX_CHUNK_CHARS = 1800


@dataclass(frozen=True)
class IngestionOptions:
    input_dir: Path
    namespace: str
    knowledge_base_id: str
    out_dir: Path
    config: AppConfig
    qdrant_collection: str | None = None
    qdrant_namespace: str | None = None
    batch_size: int = 16
    resume: bool = False


def run_knowledge_ingestion(options: IngestionOptions) -> dict[str, Any]:
    started = time.time()
    out_dir = options.out_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    namespace = display_text(options.qdrant_namespace or options.namespace)
    if not namespace:
        raise RuntimeError("namespace is required")
    records, skipped = build_ingestion_records(
        options.input_dir,
        namespace=namespace,
        knowledge_base_id=options.knowledge_base_id,
        image_output_dir=out_dir / "evidence_images",
    )
    if not records:
        raise RuntimeError(f"no supported text records found in {options.input_dir}")

    manifest_path = out_dir / "ingestion_manifest.jsonl"
    registry_rows = proof_attachment_registry_from_records(records)
    write_jsonl(manifest_path, records)
    if registry_rows:
        write_jsonl(out_dir / "proof_attachment_registry.jsonl", registry_rows)
    if skipped:
        write_jsonl(out_dir / "skipped_files.jsonl", skipped)

    client = EmbeddingClient(
        endpoint=options.config.services.embedding_endpoint,
        model=options.config.services.embedding_model,
        timeout_seconds=options.config.services.timeout_seconds,
        purpose="ingestion_embedding",
    )
    collection_name = options.qdrant_collection or options.config.qdrant.collection_name
    qdrant_path = options.config.paths.qdrant_path
    qdrant = build_qdrant_client(
        qdrant_path=qdrant_path,
        qdrant_url=options.config.qdrant.url,
        api_key_env=options.config.qdrant.api_key_env,
        prefer_grpc=options.config.qdrant.prefer_grpc,
        timeout=options.config.qdrant.timeout,
    )
    try:
        upserted, dimension = upsert_records(
            qdrant=qdrant,
            collection_name=collection_name,
            records=records,
            embedder=client,
            batch_size=options.batch_size,
            namespace=namespace,
        )
    finally:
        qdrant.close()

    summary = build_summary(
        records=records,
        skipped=skipped,
        collection_name=collection_name,
        qdrant_path=qdrant_path,
        qdrant_url=options.config.qdrant.url,
        namespace=namespace,
        embedding_endpoint=options.config.services.embedding_endpoint,
        embedding_model=options.config.services.embedding_model,
        dimension=dimension,
        upserted=upserted,
        elapsed_seconds=round(time.time() - started, 3),
    )
    write_json(out_dir / "summary.json", summary)
    write_summary_markdown(out_dir / "run_summary.md", summary)
    manifest_artifacts = {
        "ingestion_manifest": "ingestion_manifest.jsonl",
        "summary": "summary.json",
        "run_summary": "run_summary.md",
    }
    if registry_rows:
        manifest_artifacts["proof_attachment_registry"] = "proof_attachment_registry.jsonl"
    if skipped:
        manifest_artifacts["skipped_files"] = "skipped_files.jsonl"
    manifest = build_run_manifest(
        summary=summary,
        namespace=namespace,
        knowledge_base_id=options.knowledge_base_id,
        artifacts=manifest_artifacts,
    )
    write_json(out_dir / "run_manifest.json", manifest)
    return summary


def build_ingestion_records(
    input_dir: Path,
    *,
    namespace: str,
    knowledge_base_id: str,
    image_output_dir: Path | None = None,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    input_dir = input_dir.resolve()
    if not input_dir.exists():
        raise RuntimeError(f"input_dir does not exist: {input_dir}")
    files = sorted(path for path in input_dir.rglob("*") if path.is_file())
    records: list[dict[str, Any]] = []
    skipped: list[dict[str, Any]] = []
    for path in files:
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            skipped.append({"path": relative_path(path, input_dir), "reason": "unsupported_suffix"})
            continue
        try:
            chunks = list(
                extract_file_chunks(
                    path,
                    input_dir,
                    namespace=namespace,
                    knowledge_base_id=knowledge_base_id,
                    image_output_dir=image_output_dir,
                )
            )
        except Exception as exc:  # pragma: no cover - defensive path reports the real parser failure to ops.
            skipped.append({"path": relative_path(path, input_dir), "reason": f"parse_failed: {exc}"})
            continue
        for chunk in chunks:
            text = display_text(chunk["text"])
            if not text:
                continue
            for part_no, part in enumerate(split_text(text, MAX_CHUNK_CHARS), 1):
                anchor = chunk["anchor"]
                if len(text) > MAX_CHUNK_CHARS:
                    anchor = f"{anchor}#part-{part_no}"
                chunk_id = stable_chunk_id(knowledge_base_id, namespace, chunk["relative_path"], anchor, part)
                records.append(
                    {
                        "chunk_id": chunk_id,
                        "point_id": str(uuid.uuid5(uuid.NAMESPACE_URL, chunk_id)),
                        "source_type": chunk["source_type"],
                        "namespace": namespace,
                        "knowledge_base_id": knowledge_base_id,
                        "corpus_layer": "fact",
                        "embedding_policy": "embed",
                        "default_index": True,
                        "rank_boost": 1.0,
                        "text_for_embedding": part,
                        "raw_text": part,
                        "file_name": path.name,
                        "relative_path": chunk["relative_path"],
                        "sheet_name": chunk.get("sheet_name"),
                        "row_index": chunk.get("row_index"),
                        "anchor": anchor,
                        "proof_attachment_ids": chunk.get("proof_attachment_ids") or [],
                        "proof_attachments": chunk.get("proof_attachments") or [],
                        "source": {
                            "knowledge_base_id": knowledge_base_id,
                            "relative_path": chunk["relative_path"],
                            "anchor": anchor,
                            "sheet_name": chunk.get("sheet_name"),
                            "row_index": chunk.get("row_index"),
                            "proof_attachments": chunk.get("proof_attachments") or [],
                        },
                    }
                )
    return records, skipped


def extract_file_chunks(
    path: Path,
    root: Path,
    *,
    namespace: str = "",
    knowledge_base_id: str = "",
    image_output_dir: Path | None = None,
) -> Iterable[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        yield from extract_xlsx_chunks(path, root, namespace=namespace, knowledge_base_id=knowledge_base_id, image_output_dir=image_output_dir)
    elif suffix == ".docx":
        yield from extract_docx_chunks(path, root)
    else:
        yield from extract_text_chunks(path, root)


def extract_xlsx_chunks(
    path: Path,
    root: Path,
    *,
    namespace: str = "",
    knowledge_base_id: str = "",
    image_output_dir: Path | None = None,
) -> Iterable[dict[str, Any]]:
    from openpyxl import load_workbook

    registry_rows = (
        materialize_xlsx_dispimg_registry(
            path,
            root=root,
            output_dir=image_output_dir,
            namespace=namespace,
            knowledge_base_id=knowledge_base_id,
        )
        if image_output_dir
        else []
    )
    attachments_by_row = registry_by_sheet_row(registry_rows)
    workbook = load_workbook(path, read_only=True, data_only=False)
    try:
        rel = relative_path(path, root)
        for sheet in workbook.worksheets:
            for row_index, row in enumerate(sheet.iter_rows(), 1):
                values = [display_text(cell.value) for cell in row if not is_dispimg_formula(cell.value)]
                values = [value for value in values if value]
                if not values:
                    continue
                row_attachments = [attachment_from_registry(item) for item in attachments_by_row.get((sheet.title, row_index), [])]
                text = f"文件：{path.name}。Sheet：{sheet.title}。行：{row_index}。内容：" + " / ".join(values)
                yield {
                    "text": text,
                    "relative_path": rel,
                    "anchor": f"{sheet.title}!row {row_index}",
                    "source_type": "uploaded_excel_row",
                    "sheet_name": sheet.title,
                    "row_index": row_index,
                    "proof_attachment_ids": [str(item["attachment_id"]) for item in row_attachments if item.get("attachment_id")],
                    "proof_attachments": row_attachments,
                }
    finally:
        workbook.close()


def extract_docx_chunks(path: Path, root: Path) -> Iterable[dict[str, str]]:
    from docx import Document

    document = Document(path)
    rel = relative_path(path, root)
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = display_text(paragraph.text)
        if text:
            yield {
                "text": f"文件：{path.name}。段落：{index}。内容：{text}",
                "relative_path": rel,
                "anchor": f"paragraph {index}",
                "source_type": "uploaded_docx_paragraph",
            }
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            values = [display_text(cell.text) for cell in row.cells]
            values = [value for value in values if value]
            if values:
                yield {
                    "text": f"文件：{path.name}。表格：{table_index}。行：{row_index}。内容：" + " / ".join(values),
                    "relative_path": rel,
                    "anchor": f"table {table_index} row {row_index}",
                    "source_type": "uploaded_docx_table_row",
                }


def extract_text_chunks(path: Path, root: Path) -> Iterable[dict[str, str]]:
    rel = relative_path(path, root)
    text = path.read_text(encoding="utf-8", errors="replace")
    for index, part in enumerate(split_text(text, MAX_CHUNK_CHARS), 1):
        part = display_text(part)
        if part:
            yield {
                "text": f"文件：{path.name}。片段：{index}。内容：{part}",
                "relative_path": rel,
                "anchor": f"text chunk {index}",
                "source_type": "uploaded_text_chunk",
            }


def is_dispimg_formula(value: Any) -> bool:
    return isinstance(value, str) and "DISPIMG(" in value


def proof_attachment_registry_from_records(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    seen: set[str] = set()
    for record in records:
        for attachment in record.get("proof_attachments") or []:
            if not isinstance(attachment, dict):
                continue
            attachment_id = str(attachment.get("attachment_id") or "")
            if not attachment_id or attachment_id in seen:
                continue
            seen.add(attachment_id)
            rows.append(
                {
                    "attachment_id": attachment_id,
                    "file_id": attachment.get("file_id") or "",
                    "knowledge_base_id": record.get("knowledge_base_id") or "",
                    "namespace": record.get("namespace") or "",
                    "file_name": record.get("file_name") or "",
                    "relative_path": attachment.get("relative_path") or record.get("relative_path") or "",
                    "source_file_path": attachment.get("source_file_path") or "",
                    "sheet_name": attachment.get("sheet_name") or "",
                    "row_index": record.get("source", {}).get("row_index") or "",
                    "source_cell": attachment.get("source_cell") or "",
                    "image_id": attachment.get("image_id") or "",
                    "media_path": attachment.get("media_path") or "",
                    "media_content_type": attachment.get("media_content_type") or "",
                    "attachment_type": attachment.get("attachment_type") or "image",
                    "mapping_status": attachment.get("mapping_status") or "",
                    "image_path": attachment.get("image_path") or "",
                }
            )
    return rows


def upsert_records(
    *,
    qdrant: QdrantClient,
    collection_name: str,
    records: list[dict[str, Any]],
    embedder: EmbeddingClient,
    batch_size: int,
    namespace: str,
) -> tuple[int, int]:
    dimension = 0
    upserted = 0
    collection_ready = qdrant.collection_exists(collection_name)
    if collection_ready:
        qdrant.delete(
            collection_name=collection_name,
            points_selector=models.FilterSelector(
                filter=models.Filter(must=[models.FieldCondition(key="namespace", match=models.MatchValue(value=namespace))])
            ),
            wait=True,
        )
    for batch in batches(records, max(1, batch_size)):
        vectors = embedder.embed([record["text_for_embedding"] for record in batch])
        if not vectors:
            continue
        if dimension == 0:
            dimension = len(vectors[0])
        if not collection_ready:
            qdrant.create_collection(
                collection_name=collection_name,
                vectors_config=models.VectorParams(size=dimension, distance=models.Distance.COSINE),
            )
            collection_ready = True
        points = []
        for record, vector in zip(batch, vectors, strict=True):
            if len(vector) != dimension:
                raise RuntimeError(f"embedding dimension changed from {dimension} to {len(vector)}")
            payload = {key: value for key, value in record.items() if key != "point_id"}
            points.append(models.PointStruct(id=record["point_id"], vector=vector, payload=payload))
        qdrant.upsert(collection_name=collection_name, points=points, wait=True)
        upserted += len(points)
    return upserted, dimension


def build_summary(
    *,
    records: list[dict[str, Any]],
    skipped: list[dict[str, Any]],
    collection_name: str,
    qdrant_path: Path,
    qdrant_url: str,
    namespace: str,
    embedding_endpoint: str,
    embedding_model: str,
    dimension: int,
    upserted: int,
    elapsed_seconds: float,
) -> dict[str, Any]:
    return {
        "status": "completed",
        "engine": "gongkan_knowledge_ingestion",
        "collection_name": collection_name,
        "qdrant_path": str(qdrant_path),
        "qdrant_url": qdrant_url,
        "namespace": namespace,
        "record_count": len(records),
        "upserted_count": upserted,
        "skipped_file_count": len(skipped),
        "dimension": dimension,
        "embedding_endpoint": embedding_endpoint,
        "embedding_model": embedding_model,
        "elapsed_seconds": elapsed_seconds,
        "image_proof_count": sum(len(record.get("proof_attachments") or []) for record in records),
        "counts_by_source_type": dict(Counter(record["source_type"] for record in records)),
        "counts_by_corpus_layer": dict(Counter(record["corpus_layer"] for record in records)),
        "counts_by_file": dict(Counter(record["relative_path"] for record in records)),
    }


def build_run_manifest(*, summary: dict[str, Any], namespace: str, knowledge_base_id: str, artifacts: dict[str, str | None]) -> dict[str, Any]:
    now = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
    return {
        "run_id": stable_chunk_id(knowledge_base_id, namespace, now, "ingestion"),
        "created_at": now,
        "finished_at": now,
        "status": "completed",
        "engine": "gongkan_knowledge_ingestion",
        "target_namespace": namespace,
        "room_context": "",
        "rows": "",
        "judge_enabled": False,
        "writeback_enabled": False,
        "artifacts": artifacts,
        "counts": {
            "total_fields": int(summary.get("record_count") or 0),
            "answered": 0,
            "partial_clue": 0,
            "not_found": 0,
            "conflict_unresolved": 0,
            "review_required": 0,
            "writeback_allowed": 0,
            "failed": 0,
        },
    }


def write_summary_markdown(path: Path, summary: dict[str, Any]) -> None:
    lines = [
        "# Knowledge Ingestion Summary",
        "",
        f"- status: `{summary['status']}`",
        f"- namespace: `{summary['namespace']}`",
        f"- collection: `{summary['collection_name']}`",
        f"- records: **{summary['record_count']}**",
        f"- upserted: **{summary['upserted_count']}**",
        f"- skipped files: **{summary['skipped_file_count']}**",
        f"- embedding model: `{summary['embedding_model']}`",
        "",
        "## Records By Source Type",
        "",
    ]
    for key, value in sorted(summary["counts_by_source_type"].items()):
        lines.append(f"- `{key}`: {value}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def split_text(text: str, max_chars: int) -> list[str]:
    text = display_text(text)
    if len(text) <= max_chars:
        return [text] if text else []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        chunks.append(text[start:end])
        start = end
    return chunks


def batches(records: list[dict[str, Any]], batch_size: int) -> Iterable[list[dict[str, Any]]]:
    for index in range(0, len(records), batch_size):
        yield records[index : index + batch_size]


def stable_chunk_id(*parts: Any) -> str:
    text = "|".join(display_text(part) for part in parts)
    return "kb_" + hashlib.sha256(text.encode("utf-8")).hexdigest()[:24]


def relative_path(path: Path, root: Path) -> str:
    try:
        return str(path.resolve().relative_to(root.resolve()))
    except ValueError:
        return str(path)


def dumps_summary(summary: dict[str, Any]) -> str:
    return json.dumps(
        {
            "status": summary["status"],
            "namespace": summary["namespace"],
            "collection_name": summary["collection_name"],
            "record_count": summary["record_count"],
            "upserted_count": summary["upserted_count"],
            "outcome": "ready_for_retrieval",
        },
        ensure_ascii=False,
    )
