from __future__ import annotations

import argparse
import json
import mimetypes
import posixpath
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document


DEFAULT_IN = Path(__file__).resolve().parents[1] / "artifacts/03_format_probe/probed_manifest.jsonl"
DEFAULT_OUT_DIR = Path(__file__).resolve().parents[1] / "artifacts/04a_structure_parse"

NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_DRAWING = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"

DISPIMG_RE = re.compile(r'DISPIMG\("([^"]+)"', re.IGNORECASE)
CELL_RE = re.compile(r"^([A-Z]+)([0-9]+)$")
SPID_RE = re.compile(r"_x0000_s(\d+)")


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def write_jsonl(path: Path, records: list[dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8") as f:
        for record in records:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")


def write_json(path: Path, value: Any) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def tag(local: str, namespace: str = NS_MAIN) -> str:
    return f"{{{namespace}}}{local}"


def attr_rel_id() -> str:
    return f"{{{NS_REL}}}id"


def local_name(name: str) -> str:
    return name.rsplit("}", 1)[-1] if "}" in name else name


def col_to_num(col: str) -> int:
    num = 0
    for ch in col:
        num = num * 26 + ord(ch) - ord("A") + 1
    return num


def num_to_col(num: int) -> str:
    chars: list[str] = []
    while num:
        num, rem = divmod(num - 1, 26)
        chars.append(chr(ord("A") + rem))
    return "".join(reversed(chars))


def split_cell_ref(ref: str) -> tuple[int, int] | None:
    match = CELL_RE.match(ref or "")
    if not match:
        return None
    col, row = match.groups()
    return int(row), col_to_num(col)


def cell_ref(row: int, col: int) -> str:
    return f"{num_to_col(col)}{row}"


def normalize_target(base: str, target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    return posixpath.normpath(posixpath.join(posixpath.dirname(base), target))


def parse_rels(zf: zipfile.ZipFile, rels_path: str) -> dict[str, str]:
    if rels_path not in zf.namelist():
        return {}
    root = ET.fromstring(zf.read(rels_path))
    rels: dict[str, str] = {}
    for child in root:
        if local_name(child.tag) == "Relationship":
            rel_id = child.attrib.get("Id")
            target = child.attrib.get("Target")
            if rel_id and target:
                rels[rel_id] = target
    return rels


def parse_shared_strings(zf: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in zf.namelist():
        return []
    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    values: list[str] = []
    for si in root.findall(tag("si")):
        parts = [node.text or "" for node in si.iter(tag("t"))]
        values.append("".join(parts))
    return values


def parse_workbook(zf: zipfile.ZipFile) -> list[dict[str, Any]]:
    root = ET.fromstring(zf.read("xl/workbook.xml"))
    rels = parse_rels(zf, "xl/_rels/workbook.xml.rels")
    sheets: list[dict[str, Any]] = []
    for index, node in enumerate(root.findall(f".//{tag('sheet')}"), start=1):
        rel_id = node.attrib.get(attr_rel_id())
        target = rels.get(rel_id or "", "")
        path = normalize_target("xl/workbook.xml", target) if target else ""
        sheets.append(
            {
                "sheet_index": index,
                "sheet_id": node.attrib.get("sheetId"),
                "sheet_name": node.attrib.get("name", f"Sheet{index}"),
                "state": node.attrib.get("state", "visible"),
                "relationship_id": rel_id,
                "worksheet_path": path,
            }
        )
    return sheets


def parse_cell_value(cell: ET.Element, shared_strings: list[str]) -> tuple[str, Any, str | None, str | None]:
    raw_type = cell.attrib.get("t", "n")
    formula_node = cell.find(tag("f"))
    value_node = cell.find(tag("v"))

    formula_text = formula_node.text if formula_node is not None else None
    if formula_text is not None:
        cached_value = value_node.text if value_node is not None else None
        formula_kind = "wps_dispimg" if DISPIMG_RE.search(formula_text) else "formula"
        return "formula", cached_value, formula_text, formula_kind

    if raw_type == "s" and value_node is not None:
        try:
            idx = int(value_node.text or "0")
            return "shared_string", shared_strings[idx], None, None
        except (ValueError, IndexError):
            return "shared_string", value_node.text, None, None

    if raw_type == "inlineStr":
        text = "".join(node.text or "" for node in cell.iter(tag("t")))
        return "inline_string", text, None, None

    if raw_type == "str":
        return "string", value_node.text if value_node is not None else None, None, None

    if raw_type == "b":
        return "boolean", value_node.text if value_node is not None else None, None, None

    if raw_type == "e":
        return "error", value_node.text if value_node is not None else None, None, None

    return "number", value_node.text if value_node is not None else None, None, None


def parse_hidden_cols(root: ET.Element) -> set[int]:
    hidden: set[int] = set()
    cols_node = root.find(tag("cols"))
    if cols_node is None:
        return hidden
    for col in cols_node.findall(tag("col")):
        if col.attrib.get("hidden") != "1":
            continue
        try:
            min_col = int(float(col.attrib.get("min", "0")))
            max_col = int(float(col.attrib.get("max", "0")))
        except ValueError:
            continue
        hidden.update(range(min_col, max_col + 1))
    return hidden


def expand_range(range_ref: str, limit: int = 25) -> list[str]:
    if ":" not in range_ref:
        return [range_ref]
    start, end = range_ref.split(":", 1)
    start_rc = split_cell_ref(start)
    end_rc = split_cell_ref(end)
    if not start_rc or not end_rc:
        return []
    start_row, start_col = start_rc
    end_row, end_col = end_rc
    refs: list[str] = []
    for row in range(start_row, end_row + 1):
        for col in range(start_col, end_col + 1):
            refs.append(cell_ref(row, col))
            if len(refs) >= limit:
                return refs
    return refs


def parse_merges(root: ET.Element, cells_by_ref: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    merges_node = root.find(tag("mergeCells"))
    if merges_node is None:
        return []
    merges: list[dict[str, Any]] = []
    for merge in merges_node.findall(tag("mergeCell")):
        range_ref = merge.attrib.get("ref", "")
        master_cell = range_ref.split(":", 1)[0]
        merges.append(
            {
                "range": range_ref,
                "master_cell": master_cell,
                "master_value": cells_by_ref.get(master_cell, {}).get("value"),
                "covered_cells_sample": expand_range(range_ref),
            }
        )
    return merges


def parse_cellimages(zf: zipfile.ZipFile) -> dict[str, dict[str, str]]:
    names = zf.namelist()
    cellimages_path = next((name for name in names if name.lower() == "xl/cellimages.xml"), None)
    rels_path = next((name for name in names if name.lower() == "xl/_rels/cellimages.xml.rels"), None)
    if not cellimages_path or not rels_path:
        return {}

    rels = parse_rels(zf, rels_path)
    root = ET.fromstring(zf.read(cellimages_path))
    result: dict[str, dict[str, str]] = {}
    for cell_image in root.iter():
        if local_name(cell_image.tag).lower() != "cellimage":
            continue
        image_id = None
        rel_id = None
        for node in cell_image.iter():
            node_name = local_name(node.tag).lower()
            if node_name == "cnvpr":
                image_id = node.attrib.get("name")
            if node_name == "blip":
                rel_id = node.attrib.get(attr_rel_id())
                if rel_id is None:
                    rel_id = next((value for key, value in node.attrib.items() if key.lower().endswith("embed")), None)
        if not image_id or not rel_id:
            continue
        target = rels.get(rel_id)
        if not target:
            continue
        media_path = normalize_target("xl/cellimages.xml", target)
        result[image_id] = {"relationship_id": rel_id, "media_path": media_path}
    return result


def media_content_type(media_path: str) -> str | None:
    guessed, _ = mimetypes.guess_type(media_path)
    return guessed


def parse_drawing_attachments(
    zf: zipfile.ZipFile,
    worksheet_path: str,
    root: ET.Element,
    file_id: str,
    sheet_name: str,
) -> list[dict[str, Any]]:
    attachments: list[dict[str, Any]] = []
    sheet_rels_path = posixpath.join(
        posixpath.dirname(worksheet_path),
        "_rels",
        posixpath.basename(worksheet_path) + ".rels",
    )
    sheet_rels = parse_rels(zf, sheet_rels_path)
    drawing_nodes = [node for node in root.iter() if local_name(node.tag) == "drawing"]
    for drawing_node in drawing_nodes:
        drawing_rel_id = drawing_node.attrib.get(attr_rel_id())
        drawing_target = sheet_rels.get(drawing_rel_id or "")
        if not drawing_target:
            continue
        drawing_path = normalize_target(worksheet_path, drawing_target)
        if drawing_path not in zf.namelist():
            continue
        drawing_rels_path = posixpath.join(
            posixpath.dirname(drawing_path),
            "_rels",
            posixpath.basename(drawing_path) + ".rels",
        )
        drawing_rels = parse_rels(zf, drawing_rels_path)
        drawing_root = ET.fromstring(zf.read(drawing_path))
        for index, anchor in enumerate(drawing_root):
            if local_name(anchor.tag) not in {"twoCellAnchor", "oneCellAnchor", "absoluteAnchor"}:
                continue
            row = None
            col = None
            from_node = next((node for node in anchor if local_name(node.tag) == "from"), None)
            if from_node is not None:
                row_node = next((node for node in from_node if local_name(node.tag) == "row"), None)
                col_node = next((node for node in from_node if local_name(node.tag) == "col"), None)
                if row_node is not None and col_node is not None:
                    row = int(row_node.text or "0") + 1
                    col = int(col_node.text or "0") + 1
            rel_id = None
            for node in anchor.iter():
                if local_name(node.tag) == "blip":
                    rel_id = node.attrib.get(attr_rel_id())
                    break
            if not rel_id or rel_id not in drawing_rels:
                continue
            media_path = normalize_target(drawing_path, drawing_rels[rel_id])
            source_cell = cell_ref(row, col) if row and col else None
            attachments.append(
                {
                    "attachment_id": f"att_{file_id}_{sheet_name}_drawing_{index + 1}",
                    "file_id": file_id,
                    "sheet_name": sheet_name,
                    "anchor_type": "drawing",
                    "source_cell": source_cell,
                    "relationship_id": rel_id,
                    "media_path": media_path,
                    "media_content_type": media_content_type(media_path),
                    "attachment_type": "image",
                    "ocr_status": "not_required",
                    "used_for_generation": False,
                    "used_for_audit": True,
                }
            )
    return attachments


def fallback_object_content_type(prog_id: str | None, object_path: str) -> str:
    guessed = media_content_type(object_path)
    if guessed and guessed != "application/octet-stream":
        return guessed
    if prog_id == "Package":
        return "application/x-ole-package"
    if prog_id:
        return "application/vnd.openxmlformats-officedocument.oleObject"
    return "application/octet-stream"


def attr_by_local_name(node: ET.Element, name: str) -> str | None:
    for key, value in node.attrib.items():
        if local_name(key).lower() == name.lower():
            return value
    return None


def parse_vml_shape_anchors(
    zf: zipfile.ZipFile,
    worksheet_path: str,
    root: ET.Element,
    sheet_rels: dict[str, str],
) -> dict[str, dict[str, Any]]:
    legacy_node = next((node for node in root.iter() if local_name(node.tag) == "legacyDrawing"), None)
    if legacy_node is None:
        return {}
    legacy_rel_id = legacy_node.attrib.get(attr_rel_id())
    legacy_target = sheet_rels.get(legacy_rel_id or "")
    if not legacy_target:
        return {}
    vml_path = normalize_target(worksheet_path, legacy_target)
    if vml_path not in zf.namelist():
        return {}

    vml_rels_path = posixpath.join(
        posixpath.dirname(vml_path),
        "_rels",
        posixpath.basename(vml_path) + ".rels",
    )
    vml_rels = parse_rels(zf, vml_rels_path)
    vml_root = ET.fromstring(zf.read(vml_path))
    anchors: dict[str, dict[str, Any]] = {}

    for shape in vml_root.iter():
        if local_name(shape.tag) != "shape":
            continue
        spid = attr_by_local_name(shape, "spid") or attr_by_local_name(shape, "id") or ""
        match = SPID_RE.search(spid)
        if not match:
            continue
        shape_id = match.group(1)
        anchor_node = next((node for node in shape.iter() if local_name(node.tag) == "Anchor"), None)
        source_cell = None
        anchor_values: list[int] = []
        if anchor_node is not None and anchor_node.text:
            try:
                anchor_values = [int(part.strip()) for part in anchor_node.text.split(",")]
            except ValueError:
                anchor_values = []
            if len(anchor_values) >= 3:
                source_cell = cell_ref(anchor_values[2] + 1, anchor_values[0] + 1)

        image_node = next((node for node in shape.iter() if local_name(node.tag) == "imagedata"), None)
        preview_rel_id = attr_by_local_name(image_node, "relid") if image_node is not None else None
        preview_media_path = normalize_target(vml_path, vml_rels[preview_rel_id]) if preview_rel_id in vml_rels else None
        anchors[shape_id] = {
            "source_cell": source_cell,
            "anchor_values": anchor_values,
            "vml_shape_id": shape_id,
            "vml_spid": spid,
            "preview_relationship_id": preview_rel_id,
            "preview_media_path": preview_media_path,
            "preview_media_content_type": media_content_type(preview_media_path) if preview_media_path else None,
        }

    return anchors


def parse_inline_object_anchor(ole_node: ET.Element) -> dict[str, Any]:
    anchor_node = next((node for node in ole_node.iter() if local_name(node.tag) == "anchor"), None)
    if anchor_node is None:
        return {}
    from_node = next((node for node in anchor_node if local_name(node.tag) == "from"), None)
    if from_node is None:
        return {}
    row = None
    col = None
    for node in from_node:
        if local_name(node.tag) == "row" and node.text is not None:
            row = int(node.text) + 1
        elif local_name(node.tag) == "col" and node.text is not None:
            col = int(node.text) + 1
    if row is None or col is None:
        return {}
    return {
        "source_cell": cell_ref(row, col),
        "inline_anchor_from": {"row": row, "col": col},
    }


def parse_ole_attachments(
    zf: zipfile.ZipFile,
    worksheet_path: str,
    root: ET.Element,
    file_id: str,
    sheet_name: str,
) -> list[dict[str, Any]]:
    sheet_rels_path = posixpath.join(
        posixpath.dirname(worksheet_path),
        "_rels",
        posixpath.basename(worksheet_path) + ".rels",
    )
    sheet_rels = parse_rels(zf, sheet_rels_path)
    shape_anchors = parse_vml_shape_anchors(zf, worksheet_path, root, sheet_rels)

    attachments: list[dict[str, Any]] = []
    seen: set[tuple[str | None, str | None]] = set()
    for ole_node in root.iter():
        if local_name(ole_node.tag) != "oleObject":
            continue
        shape_id = ole_node.attrib.get("shapeId")
        rel_id = ole_node.attrib.get(attr_rel_id())
        key = (shape_id, rel_id)
        if key in seen:
            continue
        seen.add(key)
        target = sheet_rels.get(rel_id or "")
        if not target:
            continue
        object_path = normalize_target(worksheet_path, target)
        inline_anchor = parse_inline_object_anchor(ole_node)
        anchor = {**shape_anchors.get(shape_id or "", {}), **inline_anchor}
        prog_id = ole_node.attrib.get("progId")
        attachments.append(
            {
                "attachment_id": f"att_{file_id}_{shape_id or rel_id}_ole",
                "file_id": file_id,
                "sheet_name": sheet_name,
                "anchor_type": "ole_object",
                "source_cell": anchor.get("source_cell"),
                "relationship_id": rel_id,
                "media_path": object_path,
                "media_content_type": fallback_object_content_type(prog_id, object_path),
                "attachment_type": "embedded_object",
                "prog_id": prog_id,
                "object_shape_id": shape_id,
                "preview_relationship_id": anchor.get("preview_relationship_id"),
                "preview_media_path": anchor.get("preview_media_path"),
                "preview_media_content_type": anchor.get("preview_media_content_type"),
                "ocr_status": "not_required",
                "used_for_generation": False,
                "used_for_audit": True,
                "mapping_status": "mapped" if anchor.get("source_cell") else "unmapped",
            }
        )
    return attachments


def parse_worksheet(
    zf: zipfile.ZipFile,
    record: dict[str, Any],
    sheet: dict[str, Any],
    shared_strings: list[str],
    cellimage_map: dict[str, dict[str, str]],
) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]]]:
    worksheet_path = sheet["worksheet_path"]
    root = ET.fromstring(zf.read(worksheet_path))
    hidden_cols = parse_hidden_cols(root)
    hidden_rows: set[int] = set()
    cells: list[dict[str, Any]] = []
    cells_by_ref: dict[str, dict[str, Any]] = {}
    attachments: list[dict[str, Any]] = []
    coords: list[tuple[int, int]] = []

    for row_node in root.findall(f".//{tag('row')}"):
        try:
            row_index = int(row_node.attrib.get("r", "0"))
        except ValueError:
            row_index = 0
        row_hidden = row_node.attrib.get("hidden") == "1"
        if row_hidden and row_index:
            hidden_rows.add(row_index)
        for cell_node in row_node.findall(tag("c")):
            ref = cell_node.attrib.get("r")
            rc = split_cell_ref(ref or "")
            if not ref or not rc:
                continue
            row, col = rc
            raw_type, value, formula_text, formula_kind = parse_cell_value(cell_node, shared_strings)
            style_id = cell_node.attrib.get("s")
            cell_record: dict[str, Any] = {
                "file_id": record["file_id"],
                "sheet_index": sheet["sheet_index"],
                "sheet_name": sheet["sheet_name"],
                "cell_ref": ref,
                "row": row,
                "col": col,
                "raw_type": raw_type,
                "value": value,
                "formula_text": formula_text,
                "formula_kind": formula_kind,
                "style_id": int(style_id) if style_id is not None and style_id.isdigit() else style_id,
                "is_hidden_row": row in hidden_rows,
                "is_hidden_col": col in hidden_cols,
                "source_anchor": {
                    "file_name": record["file_name"],
                    "sheet_name": sheet["sheet_name"],
                    "cell": ref,
                },
            }
            cells.append(cell_record)
            cells_by_ref[ref] = cell_record
            coords.append((row, col))

            if formula_text:
                match = DISPIMG_RE.search(formula_text)
                if match:
                    image_id = match.group(1)
                    image_info = cellimage_map.get(image_id)
                    attachments.append(
                        {
                            "attachment_id": f"att_{record['file_id']}_{ref}_dispimg",
                            "file_id": record["file_id"],
                            "sheet_name": sheet["sheet_name"],
                            "anchor_type": "cell_formula",
                            "source_cell": ref,
                            "image_id": image_id,
                            "relationship_id": image_info.get("relationship_id") if image_info else None,
                            "media_path": image_info.get("media_path") if image_info else None,
                            "media_content_type": media_content_type(image_info["media_path"]) if image_info else None,
                            "attachment_type": "image",
                            "ocr_status": "not_required",
                            "used_for_generation": False,
                            "used_for_audit": True,
                            "mapping_status": "mapped" if image_info else "unmapped",
                        }
                    )

    merges = parse_merges(root, cells_by_ref)
    attachments.extend(parse_drawing_attachments(zf, worksheet_path, root, record["file_id"], sheet["sheet_name"]))
    attachments.extend(parse_ole_attachments(zf, worksheet_path, root, record["file_id"], sheet["sheet_name"]))

    dimension_node = root.find(tag("dimension"))
    declared_dimension = dimension_node.attrib.get("ref") if dimension_node is not None else None
    min_row = min((row for row, _ in coords), default=None)
    max_row = max((row for row, _ in coords), default=None)
    min_col = min((col for _, col in coords), default=None)
    max_col = max((col for _, col in coords), default=None)
    actual_min_cell = cell_ref(min_row, min_col) if min_row and min_col else None
    actual_max_cell = cell_ref(max_row, max_col) if max_row and max_col else None
    formula_count = sum(1 for cell in cells if cell["formula_text"])
    dispimg_count = sum(1 for cell in cells if cell["formula_kind"] == "wps_dispimg")

    sheet_summary = {
        **sheet,
        "declared_dimension": declared_dimension,
        "actual_min_cell": actual_min_cell,
        "actual_max_cell": actual_max_cell,
        "actual_dimension": f"{actual_min_cell}:{actual_max_cell}" if actual_min_cell and actual_max_cell else None,
        "non_empty_cell_count": len(cells),
        "merge_count": len(merges),
        "formula_count": formula_count,
        "dispimg_formula_count": dispimg_count,
        "hidden_row_count": len(hidden_rows),
        "hidden_col_count": len(hidden_cols),
        "attachment_count": len(attachments),
        "merges": merges,
        "parse_status": "ok",
    }
    return sheet_summary, cells, attachments


def parse_xlsx(record: dict[str, Any], out_dir: Path) -> dict[str, Any]:
    source_path = Path(record["source_path"])
    workbook_dir = out_dir / "workbooks"
    worksheet_dir = out_dir / "worksheets"
    attachment_dir = out_dir / "attachments"
    diagnostics_dir = out_dir / "diagnostics"
    for directory in [workbook_dir, worksheet_dir, attachment_dir, diagnostics_dir]:
        directory.mkdir(parents=True, exist_ok=True)

    all_attachments: list[dict[str, Any]] = []
    sheet_summaries: list[dict[str, Any]] = []
    diagnostics: list[str] = []

    with zipfile.ZipFile(source_path) as zf:
        shared_strings = parse_shared_strings(zf)
        sheets = parse_workbook(zf)
        cellimage_map = parse_cellimages(zf)
        for sheet in sheets:
            if sheet["worksheet_path"] not in zf.namelist():
                diagnostics.append(f"worksheet missing: {sheet['worksheet_path']}")
                continue
            sheet_summary, cells, attachments = parse_worksheet(zf, record, sheet, shared_strings, cellimage_map)
            sheet_summaries.append(sheet_summary)
            all_attachments.extend(attachments)
            write_jsonl(
                worksheet_dir / f"{record['file_id']}.{sheet['sheet_index']:02d}.cells.jsonl",
                cells,
            )
            write_json(
                worksheet_dir / f"{record['file_id']}.{sheet['sheet_index']:02d}.sheet.json",
                sheet_summary,
            )

    workbook_record = {
        "file_id": record["file_id"],
        "file_name": record["file_name"],
        "relative_path": record["relative_path"],
        "document_role": record["document_role"],
        "data_center_id": record.get("data_center_id"),
        "parser_type": record["parser_type"],
        "parse_status": "ok",
        "sheet_count": len(sheet_summaries),
        "sheets": sheet_summaries,
        "total_cell_count": sum(sheet["non_empty_cell_count"] for sheet in sheet_summaries),
        "total_formula_count": sum(sheet["formula_count"] for sheet in sheet_summaries),
        "total_dispimg_formula_count": sum(sheet["dispimg_formula_count"] for sheet in sheet_summaries),
        "total_attachment_count": len(all_attachments),
    }
    write_json(workbook_dir / f"{record['file_id']}.workbook.json", workbook_record)
    write_jsonl(attachment_dir / f"{record['file_id']}.attachments.jsonl", all_attachments)
    write_json(diagnostics_dir / f"{record['file_id']}.diagnostics.json", {"diagnostics": diagnostics})
    return workbook_record


def parse_docx(record: dict[str, Any], out_dir: Path) -> dict[str, Any]:
    source_path = Path(record["source_path"])
    documents_dir = out_dir / "documents"
    attachment_dir = out_dir / "attachments"
    diagnostics_dir = out_dir / "diagnostics"
    for directory in [documents_dir, attachment_dir, diagnostics_dir]:
        directory.mkdir(parents=True, exist_ok=True)

    doc = Document(source_path)
    blocks: list[dict[str, Any]] = []
    block_index = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        block_index += 1
        blocks.append(
            {
                "file_id": record["file_id"],
                "block_type": "paragraph",
                "block_index": block_index,
                "text": text,
                "style_name": paragraph.style.name if paragraph.style else None,
                "source_anchor": {"file_name": record["file_name"], "block_index": block_index},
            }
        )

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        cells: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows, start=1):
            values: list[str] = []
            for col_index, cell in enumerate(row.cells, start=1):
                text = cell.text.strip()
                values.append(text)
                cells.append(
                    {
                        "file_id": record["file_id"],
                        "block_type": "table_cell",
                        "table_index": table_index,
                        "row_index": row_index,
                        "col_index": col_index,
                        "text": text,
                        "source_anchor": {
                            "file_name": record["file_name"],
                            "table_index": table_index,
                            "row_index": row_index,
                            "col_index": col_index,
                        },
                    }
                )
            rows.append(values)
        tables.append({"table_index": table_index, "row_count": len(rows), "rows_sample": rows[:5], "cells": cells})

    attachments: list[dict[str, Any]] = []
    with zipfile.ZipFile(source_path) as zf:
        media_names = [
            n
            for n in zf.namelist()
            if n.startswith("word/media/") and not n.endswith("/") and posixpath.basename(n)
        ]
        for idx, name in enumerate(media_names, start=1):
            attachments.append(
                {
                    "attachment_id": f"att_{record['file_id']}_word_media_{idx}",
                    "file_id": record["file_id"],
                    "anchor_type": "docx_media",
                    "media_path": name,
                    "media_content_type": media_content_type(name),
                    "attachment_type": "image",
                    "ocr_status": "not_required",
                    "used_for_generation": False,
                    "used_for_audit": True,
                }
            )
        embedding_count = sum(1 for n in zf.namelist() if n.startswith("word/embeddings/"))

    doc_record = {
        "file_id": record["file_id"],
        "file_name": record["file_name"],
        "relative_path": record["relative_path"],
        "document_role": record["document_role"],
        "data_center_id": record.get("data_center_id"),
        "parser_type": record["parser_type"],
        "parse_status": "ok",
        "paragraph_count": len(blocks),
        "table_count": len(tables),
        "attachment_count": len(attachments),
        "embedding_count": embedding_count,
        "paragraphs": blocks,
        "tables": tables,
    }
    write_json(documents_dir / f"{record['file_id']}.docx.json", doc_record)
    write_jsonl(attachment_dir / f"{record['file_id']}.attachments.jsonl", attachments)
    write_json(diagnostics_dir / f"{record['file_id']}.diagnostics.json", {"diagnostics": []})
    return doc_record


def write_visualization(out_dir: Path, file_results: list[dict[str, Any]]) -> None:
    status_counts = Counter(result["parse_status"] for result in file_results)
    parser_counts = Counter(result["parser_type"] for result in file_results)
    lines: list[str] = []
    lines.append("# Step 04A 确定性结构解析可视化\n")
    lines.append(f"- 输入文件数：**{len(file_results)}**")
    lines.append("- 本步骤只做结构解析，不做语义判断，不做 OCR。\n")

    lines.append("## 状态统计\n")
    lines.append("| parse_status | count |")
    lines.append("|---|---:|")
    for status, count in sorted(status_counts.items()):
        lines.append(f"| `{status}` | {count} |")

    lines.append("\n## parser_type 统计\n")
    lines.append("| parser_type | count |")
    lines.append("|---|---:|")
    for parser_type, count in sorted(parser_counts.items()):
        lines.append(f"| `{parser_type}` | {count} |")

    lines.append("\n## 文件结构摘要\n")
    lines.append("| file | parser | status | sheets/tables | cells/paras | formulas | attachments | note |")
    lines.append("|---|---|---|---:|---:|---:|---:|---|")
    for result in file_results:
        if result["parser_type"] == "xlsx_ooxml":
            lines.append(
                f"| `{result['relative_path']}` | `{result['parser_type']}` | `{result['parse_status']}` | "
                f"{result.get('sheet_count', 0)} | {result.get('total_cell_count', 0)} | "
                f"{result.get('total_formula_count', 0)} | {result.get('total_attachment_count', 0)} |  |"
            )
        elif result["parser_type"] == "docx_ooxml":
            lines.append(
                f"| `{result['relative_path']}` | `{result['parser_type']}` | `{result['parse_status']}` | "
                f"{result.get('table_count', 0)} | {result.get('paragraph_count', 0)} | 0 | "
                f"{result.get('attachment_count', 0)} | docx |"
            )
        else:
            lines.append(
                f"| `{result['relative_path']}` | `{result['parser_type']}` | `{result['parse_status']}` | "
                f"0 | 0 | 0 | 0 | `{result.get('fallback_action') or ''}` |"
            )

    lines.append("\n## Sheet 样例\n")
    lines.append("| file | sheet | actual_dimension | cells | formulas | dispimg | attachments |")
    lines.append("|---|---|---|---:|---:|---:|---:|")
    for result in file_results:
        for sheet in result.get("sheets", [])[:3]:
            lines.append(
                f"| `{result['relative_path']}` | `{sheet['sheet_name']}` | `{sheet.get('actual_dimension')}` | "
                f"{sheet.get('non_empty_cell_count', 0)} | {sheet.get('formula_count', 0)} | "
                f"{sheet.get('dispimg_formula_count', 0)} | {sheet.get('attachment_count', 0)} |"
            )

    (out_dir / "visualization.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def parse_one(record: dict[str, Any], out_dir: Path) -> dict[str, Any]:
    if record["parser_type"] == "xlsx_ooxml" and record["parse_status"] == "ok":
        return parse_xlsx(record, out_dir)
    if record["parser_type"] == "docx_ooxml" and record["parse_status"] == "ok":
        return parse_docx(record, out_dir)

    diagnostics_dir = out_dir / "diagnostics"
    diagnostics_dir.mkdir(parents=True, exist_ok=True)
    result = {
        "file_id": record["file_id"],
        "file_name": record["file_name"],
        "relative_path": record["relative_path"],
        "document_role": record["document_role"],
        "data_center_id": record.get("data_center_id"),
        "parser_type": record["parser_type"],
        "parse_status": record["parse_status"],
        "fallback_action": record.get("fallback_action"),
        "diagnostics": record.get("diagnostics", []),
    }
    write_json(diagnostics_dir / f"{record['file_id']}.diagnostics.json", result)
    return result


def run(input_path: Path = DEFAULT_IN, out_dir: Path = DEFAULT_OUT_DIR) -> list[dict[str, Any]]:
    for subdir in ["workbooks", "worksheets", "documents", "attachments", "diagnostics"]:
        (out_dir / subdir).mkdir(parents=True, exist_ok=True)

    records = read_jsonl(input_path)
    file_results = [parse_one(record, out_dir) for record in records]
    write_jsonl(out_dir / "files.jsonl", file_results)
    report = {
        "total_files": len(file_results),
        "parse_status_counts": dict(Counter(result["parse_status"] for result in file_results)),
        "parser_type_counts": dict(Counter(result["parser_type"] for result in file_results)),
        "total_xlsx_cells": sum(result.get("total_cell_count", 0) for result in file_results),
        "total_xlsx_formulas": sum(result.get("total_formula_count", 0) for result in file_results),
        "total_xlsx_dispimg_formulas": sum(result.get("total_dispimg_formula_count", 0) for result in file_results),
        "total_attachments": sum(
            result.get("total_attachment_count", result.get("attachment_count", 0)) for result in file_results
        ),
        "files": file_results,
    }
    write_json(out_dir / "parse_report.json", report)
    write_visualization(out_dir, file_results)
    return file_results


def main() -> None:
    parser = argparse.ArgumentParser(description="Step 04A: deterministic structure parsing.")
    parser.add_argument("--input", type=Path, default=DEFAULT_IN)
    parser.add_argument("--out-dir", type=Path, default=DEFAULT_OUT_DIR)
    args = parser.parse_args()
    results = run(args.input, args.out_dir)
    print(f"parsed structure for {len(results)} files -> {args.out_dir}")


if __name__ == "__main__":
    main()
